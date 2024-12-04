import os
import mimetypes
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
import xml.etree.ElementTree as ET
import spacy
from spacy.cli import download
from django.http import FileResponse
from trtokenizer import SentenceTokenizer
from ..models import Memory
from ..models import MemoryAsset
from ..models import Settings
import json
import openai
from django.http import JsonResponse

# ABBYY Cloud OCR credentials
application_id = '88bff69a-a1ab-453c-9b8c-1dfdd03de30c'
password = 'fVXQz2zQlLRktS7gLX9ZiIWo'
base_url = 'https://cloud-westus.ocrsdk.com'
# Function to detect if the file is PDF or an image

LANGUAGE_CODES = {
    "French": "fr",
    "Arabic": "ar",
    "Turkish": "tr",
    "English": "en",
}

def load_spacy_model(model_name="en_core_web_sm"):
    try:
        # Try loading the model
        nlp = spacy.load(model_name)
    except OSError:
        # If model is not found, download it
        print(f"Model '{model_name}' not found. Downloading...")
        download(model_name)
        # Load the model after downloading
        nlp = spacy.load(model_name)
    return nlp

# Loading spacy models
nlp = load_spacy_model("en_core_web_sm")
sentence_tokenizer = SentenceTokenizer()  # Initialize trtokenizer for Turkish


def detect_file_type(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type == "application/pdf":
        return "pdf"
    elif mime_type and mime_type.startswith("image"):
        return "image"
    else:
        return None

# Function to parse XML responses and extract taskId, status, and resultUrl
def parse_xml_response(response_text):
    root = ET.fromstring(response_text)
    task_element = root.find('task')
    if task_element is not None:
        return {
            'taskId': task_element.attrib.get('id'),
            'status': task_element.attrib.get('status'),
            'resultUrl': task_element.attrib.get('resultUrl'),
            'estimatedProcessingTime': task_element.attrib.get('estimatedProcessingTime')
        }
    return None

# Submit file to ABBYY OCR
def submit_file_for_ocr(file, file_type, language):
    url = f'{base_url}/processImage'
    auth = (application_id, password)
    files = {'file': file}
    data = {
        'language': language,
        'exportFormat': 'docx',
        'textType': 'normal,handprinted,gothic,typewriter,cmc7',
        'correctSkew': 'true',
        'correctOrientation': 'true',
        'imageSource': 'auto'
    }
    response = requests.post(url, files=files, auth=auth, data=data)
    response.raise_for_status()
    return parse_xml_response(response.text)

# Function to save OCR result as DOCX file
def get_ocr_result(task_id, output_docx_path):
    url = f'{base_url}/getTaskStatus'
    auth = (application_id, password)
    params = {'taskId': task_id}

    while True:
        response = requests.get(url, params=params, auth=auth)
        response.raise_for_status()
        parsed_response = parse_xml_response(response.text)

        if parsed_response is None:
            return None

        status = parsed_response.get('status', 'Unknown')

        if status == 'Completed':
            result_url = parsed_response.get('resultUrl', None)
            if result_url:
                docx_response = requests.get(result_url)
                with open(output_docx_path, 'wb') as f:
                    f.write(docx_response.content)
                return output_docx_path
        elif status == 'ProcessingFailed':
            return None
        time.sleep(5)

# Main function to submit the file, extract text, and convert to DOCX
def convert_pdf_to_docx(file, output_docx_path, source_language):
    file_type = detect_file_type(file.name)
    if not file_type:
        return

    ocr_response = submit_file_for_ocr(file, file_type, source_language)
    if ocr_response and 'taskId' in ocr_response:
        get_ocr_result(ocr_response['taskId'], output_docx_path)

def create_translated_file(input_path, output_path, target_language, source_language):
    """
    Translates a Word document's content by replacing text based on memory entries.
    """
    def replace_text_in_runs(runs, memory_dict):
        # Replace text directly using the dictionary for efficiency
        for run in runs:
            if run.text in memory_dict:
                run.text = memory_dict[run.text]

    def replace_text_in_paragraphs(paragraphs, memory_dict):
        for para in paragraphs:
            replace_text_in_runs(para.runs, memory_dict)

    def replace_text_in_tables(tables, memory_dict):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_paragraphs(cell.paragraphs, memory_dict)

    # Load the document
    doc = Document(input_path)
    
    # Pre-load memory entries into a dictionary for faster lookups
    memories = Memory.objects.filter(
        source_language=source_language,
        target_language=target_language
    ).exclude(target_text='')
    memory_dict = {memory.source_text: memory.target_text for memory in memories}

    # Replace text in paragraphs
    replace_text_in_paragraphs(doc.paragraphs, memory_dict)

    # Replace text in tables
    replace_text_in_tables(doc.tables, memory_dict)

    # Replace text in headers and footers
    for section in doc.sections:
        replace_text_in_paragraphs(section.header.paragraphs, memory_dict)
        replace_text_in_paragraphs(section.footer.paragraphs, memory_dict)

    # Save the updated document
    if os.path.exists(output_path):
        os.remove(output_path)  # Ensure no conflict with an existing file

    doc.save(output_path)

def extract_sentences_for_translation(docx_path, language="turkish"):
    doc = Document(docx_path)
    extracted_text = {}

    # Function to extract text from drawing elements
    def extract_drawing_text(element):
        drawing_texts = []
        for drawing in element.findall(".//" + qn("w:drawing")):
            for t in drawing.findall(".//" + qn("a:t")):  # 'a:t' tags hold text in drawing shapes
                if t.text:
                    drawing_texts.append(t.text.strip())
        return drawing_texts

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and text not in extracted_text:
            extracted_text[text] = ""  # Placeholder for translation

        # Extract and process any text in drawing elements within the paragraph
        drawing_texts = extract_drawing_text(para._element)
        for text in drawing_texts:
            if text and text not in extracted_text:
                extracted_text[text] = ""  # Placeholder for translation

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in extracted_text:
                    extracted_text[text] = ""  # Placeholder for translation

                # Extract and process any text in drawing elements within the cell
                drawing_texts = extract_drawing_text(cell._element)
                for text in drawing_texts:
                    if text and text not in extracted_text:
                        extracted_text[text] = ""  # Placeholder for translation

    # Extract text from headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            text = para.text.strip()
            if text and text not in extracted_text:
                extracted_text[text] = ""  # Placeholder for translation

            # Extract and process any text in drawing elements within the header paragraph
            drawing_texts = extract_drawing_text(para._element)
            for text in drawing_texts:
                if text and text not in extracted_text:
                    extracted_text[text] = ""  # Placeholder for translation

    # Extract text from footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            text = para.text.strip()
            if text and text not in extracted_text:
                extracted_text[text] = ""  # Placeholder for translation

            # Extract and process any text in drawing elements within the footer paragraph
            drawing_texts = extract_drawing_text(para._element)
            for text in drawing_texts:
                if text and text not in extracted_text:
                    extracted_text[text] = ""  # Placeholder for translation

    return extracted_text

def format_extracted_sentences(extracted_sentences, memories):
    formatted_sentences = []
    
    for idx, sentence in enumerate(extracted_sentences.keys(), start=1):
        # Find the corresponding memory entry for the sentence
        matching_memory = memories.filter(source_text=sentence).first()
        
        # Append the formatted object with an ID to the list
        formatted_sentences.append({
            "id": idx,  # Unique ID for each sentence
            "originalText": sentence,
            "translatedText": matching_memory.target_text if matching_memory else ""  # Use target_text if available
        })
    
    return formatted_sentences

# Convert PDF to DOCX API
class ConvertPDFToDocxAPI(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        file = request.FILES.get('file')
        source_language = request.data.get('sourceLanguage', 'English')

        if not file:
            return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

        file_type = detect_file_type(file.name)
        if not file_type:
            return Response({'error': 'Unsupported file type'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Submit the file to ABBYY OCR
            ocr_response = submit_file_for_ocr(file, file_type, source_language)
            print(ocr_response, 'this is the ocr response')
            if not ocr_response or 'taskId' not in ocr_response:
                return Response({'error': 'Failed to process the file'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Return taskId and estimated processing time
            return Response({
                'taskId': ocr_response['taskId'],
                'estimatedProcessingTime': ocr_response.get('estimatedProcessingTime', '5000')
            }, status=200)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# Get Task Status API
class GetTaskStatusAPI(APIView):
    def get(self, request, *args, **kwargs):
        task_id = request.query_params.get('taskId')
        source_language = request.query_params.get('source_language')
        target_language = request.query_params.get('target_language')  # Corrected typo

        if not task_id:
            return Response({'error': 'Task ID is required'}, status=status.HTTP_400_BAD_REQUEST)

        url = f'{base_url}/getTaskStatus'
        auth = (application_id, password)
        params = {'taskId': task_id}

        # Use LANGUAGE_CODES to map the languages
        memories = Memory.objects.filter(
            source_language=LANGUAGE_CODES.get(source_language, 'en'),
            target_language=LANGUAGE_CODES.get(target_language, 'fr')  # Corrected typo here
        )

        try:
            response = requests.get(url, params=params, auth=auth)
            response.raise_for_status()
            parsed_response = parse_xml_response(response.text)

            if parsed_response is None:
                return Response({'error': 'Invalid response from ABBYY'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            status = parsed_response.get('status')
            if status == 'Completed':
                # Fetch the result file and process it
                result_url = parsed_response.get('resultUrl')
                if not result_url:
                    return Response({'error': 'Result URL not found'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                output_docx_path = 'result_file.docx'
                docx_response = requests.get(result_url)
                with open(output_docx_path, 'wb') as f:
                    f.write(docx_response.content)

                # Extract text and format it for frontend
                extracted_sentences = extract_sentences_for_translation(output_docx_path)
                response_data = format_extracted_sentences(extracted_sentences, memories)
                return Response({"data": response_data, "status": status}, status=200)

            return Response({
                'taskId': parsed_response.get('taskId'),
                'status': parsed_response.get('status'),
                'estimatedProcessingTime': parsed_response.get('estimatedProcessingTime')
            }, status=200)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# API View for downloading the original DOCX document
class DownloadOriginalDocxAPI(APIView):
    def get(self, request, *args, **kwargs):
        output_docx_path = 'result_file.docx'
        
        # Check if the file exists
        if not os.path.exists(output_docx_path):
            return Response({'error': 'File not found'}, status=status.HTTP_404_NOT_FOUND)

        # Return the file as a downloadable attachment
        response = FileResponse(open(output_docx_path, 'rb'), as_attachment=True)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(output_docx_path)}"'
        return response

# API View for downloading the translated DOCX document
class DownloadReplacedDocxAPI(APIView):
    def get(self, request, *args, **kwargs):
        # Define paths and parameters
        output_docx_path = 'result_file.docx'  # Input DOCX file
        replaced_output_docx_path = 'replaced_design_original_file_name.docx'  # Output DOCX file

        # Get target and source languages as human-readable names from query params
        target_language_name = request.query_params.get('target_language', 'English')  # Default to English
        source_language_name = request.query_params.get('source_language', 'French')  # Default to French

        # Convert language names to codes using the dictionary
        target_language = LANGUAGE_CODES.get(target_language_name, 'en')  # Default to 'en' if not found
        source_language = LANGUAGE_CODES.get(source_language_name, 'fr')  # Default to 'fr' if not found

        # Ensure the input file exists before proceeding
        if not os.path.exists(output_docx_path):
            return Response({'error': 'Original file not found'}, status=status.HTTP_404_NOT_FOUND)

        # Call the create_translated_file function
        try:
            create_translated_file(output_docx_path, replaced_output_docx_path, target_language, source_language)
        except Exception as e:
            return Response({'error': f'Error generating replaced file: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Check if the replaced file was created successfully
        if not os.path.exists(replaced_output_docx_path):
            return Response({'error': 'Replaced file not found'}, status=status.HTTP_404_NOT_FOUND)

        # Return the replaced file as a downloadable attachment
        response = FileResponse(open(replaced_output_docx_path, 'rb'), as_attachment=True)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(replaced_output_docx_path)}"'
        return response

class TranslateRecordsView(APIView):
    def post(self, request):
        # Validate required parameters
        file = request.FILES.get('file')
        source_language = request.POST.get('source_language')
        target_language = request.POST.get('target_language')
        gpt_key = request.POST.get('gptKey')

        if not all([file, source_language, target_language, gpt_key]):
            return JsonResponse({"error": "Missing file or required parameters."}, status=400)

        # Set the OpenAI API key
        openai.api_key = gpt_key

        # Process the uploaded file
        try:
            file_content = file.read().decode('utf-8')
            records = json.loads(file_content)
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format in the uploaded file."}, status=400)

        # Ensure the source and target languages are valid
        source_language_code = LANGUAGE_CODES.get(source_language)
        target_language_code = LANGUAGE_CODES.get(target_language)
        if not source_language_code or not target_language_code:
            return JsonResponse({"error": "Invalid source or target language."}, status=400)

        # Prepare records for translation
        untranslated_records = [
            record.get("originalText")
            for record in records if record.get("originalText")
        ]
        if not untranslated_records:
            return JsonResponse({"error": "No valid records to translate."}, status=400)

        # Batch processing for OpenAI API
        translated_records = []
        try:
            for source_text in untranslated_records:
                response = openai.ChatCompletion.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": f"You are a professional translator. Translate the following text from {source_language} to {target_language}. If a translation isn't possible or the text doesn't require translation, return it as it is without adding any unrelated content."},
                        {"role": "user", "content": source_text}
                    ],
                    max_tokens=500,
                )
                translated_text = response.choices[0].message['content'].strip()
                translated_records.append({
                    "source_text": source_text,
                    "target_text": translated_text,
                    "gptGenerated": True  # Add this flag to indicate translation by GPT
                })

        except openai.error.OpenAIError as e:
            return JsonResponse({"error": f"OpenAI API error: {e}"}, status=500)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

        # Return translated records
        return JsonResponse({"translatedRecords": translated_records}, status=200)

class SaveApplicationSettings(APIView):
    def get(self, request, *args, **kwargs):
        """
        Retrieve the settings and return them.
        """
        try:
            settings = Settings.objects.first()
            if settings:
                data = {
                    "chat_api_key": settings.chat_api_key,
                    "abby_app_id": settings.abby_app_id,
                    "abby_password": settings.abby_password,
                }
                return JsonResponse(data, status=200)
            else:
                return JsonResponse({"message": "No settings found."}, status=404)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    def put(self, request, *args, **kwargs):
        """
        Update the settings if they exist, or create new settings if none exist.
        """
        try:
            # Parse the JSON body
            import json
            body = json.loads(request.body)
            payload = {key: value.strip() for key, value in body.items() if isinstance(value, str) and value.strip()}

            if not payload:
                return JsonResponse({"error": "No valid data provided."}, status=400)

            # Check if settings already exist
            settings = Settings.objects.first()

            if settings:
                # Update only the fields provided in the payload
                for key, value in payload.items():
                    if hasattr(settings, key):
                        setattr(settings, key, value)
                settings.save()
                return JsonResponse({"message": "Settings updated successfully."}, status=200)
            else:
                # Create a new settings record
                settings = Settings(**payload)
                settings.save()
                return JsonResponse({"message": "Settings created successfully.", settings: settings}, status=201)
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format."}, status=400)
        except Exception as e:
            print(f"{e}")
            return JsonResponse({"error": str(e)}, status=500)