import os
import mimetypes
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
import xml.etree.ElementTree as ET
import spacy
from spacy.cli import download
from django.http import FileResponse

# ABBYY Cloud OCR credentials
application_id = 'aa8da2ea-0f0b-4de8-b246-2215b67aabcb'
password = 'vEhImFfVztzs8k7mqzerFfrL'
base_url = 'https://cloud-westus.ocrsdk.com'
# Function to detect if the file is PDF or an image

def load_spacy_model(model_name="en_core_web_sm"):
    try:
        # Try loading the model
        nlp = spacy.load(model_name)
    except OSError:
        # If model is not found, download it
        print(f"Model '{model_name}' not found. Downloading...")
        download(model_name)
        # Load the model after downloading
        nlp = spacy.load(model_name)
    return nlp

# Loading spacy models
nlp = load_spacy_model("en_core_web_sm")


def detect_file_type(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type == "application/pdf":
        return "pdf"
    elif mime_type and mime_type.startswith("image"):
        return "image"
    else:
        return None

# Function to parse XML responses and extract taskId, status, and resultUrl
def parse_xml_response(response_text):
    root = ET.fromstring(response_text)
    task_element = root.find('task')
    if task_element is not None:
        return {
            'taskId': task_element.attrib.get('id'),
            'status': task_element.attrib.get('status'),
            'resultUrl': task_element.attrib.get('resultUrl')
        }
    return None

# Function to submit file to ABBYY OCR (PDF or image)
def submit_file_for_ocr(file, file_type, language):
    if file_type == "pdf" or file_type == "image":
        url = f'{base_url}/processImage'
        export_format = "docx"
    else:
        return None

    auth = (application_id, password)
    files = {'file': file}
    data = {'language': language, 'exportFormat': export_format}
    response = requests.post(url, files=files, auth=auth, data=data)
    response.raise_for_status()
    return parse_xml_response(response.text)

# Function to save OCR result as DOCX file
def get_ocr_result(task_id, output_docx_path):
    url = f'{base_url}/getTaskStatus'
    auth = (application_id, password)
    params = {'taskId': task_id}

    while True:
        response = requests.get(url, params=params, auth=auth)
        response.raise_for_status()
        parsed_response = parse_xml_response(response.text)

        if parsed_response is None:
            return None

        status = parsed_response.get('status', 'Unknown')

        if status == 'Completed':
            result_url = parsed_response.get('resultUrl', None)
            if result_url:
                docx_response = requests.get(result_url)
                with open(output_docx_path, 'wb') as f:
                    f.write(docx_response.content)
                return output_docx_path
        elif status == 'ProcessingFailed':
            return None
        time.sleep(5)

# Main function to submit the file, extract text, and convert to DOCX
def convert_pdf_to_docx(file, output_docx_path, source_language):
    file_type = detect_file_type(file.name)
    if not file_type:
        return

    ocr_response = submit_file_for_ocr(file, file_type, source_language)
    if ocr_response and 'taskId' in ocr_response:
        get_ocr_result(ocr_response['taskId'], output_docx_path)

# Function to replace all text with "Lorem ipsum" and save to new file
def create_replaced_file(input_path, replaced_output_path):
    doc = Document(input_path)
    lorem_text = "Lorem ipsum"

    # Replace text in paragraphs
    for para in doc.paragraphs:
        para.text = lorem_text

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = lorem_text

    # Replace text in headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            para.text = lorem_text

    # Replace text in footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            para.text = lorem_text

    doc.save(replaced_output_path)
    

# Function to extract all sentences from a DOCX file, including headers, footers, drawings, etc.
def extract_sentences_for_translation(docx_path):
    doc = Document(docx_path)
    extracted_sentences = {}

    # Helper function for sentence tokenization using spaCy
    def spacy_sent_tokenize(text):
        doc = nlp(text)
        return [sent.text.strip() for sent in doc.sents]  # Extract sentences from spaCy doc object

    # Function to extract text from drawing elements
    def extract_drawing_text(element):
        drawing_texts = []
        for drawing in element.findall(".//" + qn("w:drawing")):
            for t in drawing.findall(".//" + qn("a:t")):  # 'a:t' tags hold text in drawing shapes
                if t.text:
                    drawing_texts.append(t.text.strip())
        return drawing_texts

    # Extract sentences from paragraphs
    for para in doc.paragraphs:
        sentences = spacy_sent_tokenize(para.text)
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and sentence not in extracted_sentences:
                extracted_sentences[sentence] = ""  # Placeholder for translation

        # Extract and process any text in drawing elements within the paragraph
        drawing_texts = extract_drawing_text(para._element)
        for text in drawing_texts:
            sentences = spacy_sent_tokenize(text)
            for sentence in sentences:
                if sentence and sentence not in extracted_sentences:
                    extracted_sentences[sentence] = ""  # Placeholder for translation

    # Extract sentences from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                sentences = spacy_sent_tokenize(cell.text)
                for sentence in sentences:
                    sentence = sentence.strip()
                    if sentence and sentence not in extracted_sentences:
                        extracted_sentences[sentence] = ""  # Placeholder for translation

                # Extract and process any text in drawing elements within the cell
                drawing_texts = extract_drawing_text(cell._element)
                for text in drawing_texts:
                    sentences = spacy_sent_tokenize(text)
                    for sentence in sentences:
                        if sentence and sentence not in extracted_sentences:
                            extracted_sentences[sentence] = ""  # Placeholder for translation

    # Extract sentences from headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            sentences = spacy_sent_tokenize(para.text)
            for sentence in sentences:
                sentence = sentence.strip()
                if sentence and sentence not in extracted_sentences:
                    extracted_sentences[sentence] = ""  # Placeholder for translation

            # Extract and process any text in drawing elements within the header paragraph
            drawing_texts = extract_drawing_text(para._element)
            for text in drawing_texts:
                sentences = spacy_sent_tokenize(text)
                for sentence in sentences:
                    if sentence and sentence not in extracted_sentences:
                        extracted_sentences[sentence] = ""  # Placeholder for translation

    # Extract sentences from footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            sentences = spacy_sent_tokenize(para.text)
            for sentence in sentences:
                sentence = sentence.strip()
                if sentence and sentence not in extracted_sentences:
                    extracted_sentences[sentence] = ""  # Placeholder for translation

            # Extract and process any text in drawing elements within the footer paragraph
            drawing_texts = extract_drawing_text(para._element)
            for text in drawing_texts:
                sentences = spacy_sent_tokenize(text)
                for sentence in sentences:
                    if sentence and sentence not in extracted_sentences:
                        extracted_sentences[sentence] = ""  # Placeholder for translation

    return extracted_sentences

# Function to format extracted sentences as [{originalText: "", translatedText: ""}]
def format_extracted_sentences(extracted_sentences):
    formatted_sentences = []
    
    for sentence in extracted_sentences.keys():
        # Append the formatted object to the list
        formatted_sentences.append({
            "originalText": sentence,
            "translatedText": ""  # Empty string for untranslated text
        })
    
    return formatted_sentences

# Django API View to convert PDF to DOCX
class ConvertPDFToDocxAPI(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        file = request.FILES.get('file')
        source_language = request.data.get('language', 'English')

        if not file:
            return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

        output_docx_path = 'design_original_file_name.docx'
        replaced_output_docx_path = 'replaced_design_original_file_name.docx'
        print(f"source_language{source_language}")
        # Convert PDF to DOCX
        convert_pdf_to_docx(file, output_docx_path, source_language)

        # Create replaced DOCX with "Lorem ipsum" content
        create_replaced_file(output_docx_path, replaced_output_docx_path)
        extracted_sentences = extract_sentences_for_translation(output_docx_path)
        # Format extracted sentences for frontend
        response_data = format_extracted_sentences(extracted_sentences)

        # Return the formatted extracted sentences as JSON response
        return Response(response_data, status=200)

# API View for downloading the original DOCX document
class DownloadOriginalDocxAPI(APIView):
    def get(self, request, *args, **kwargs):
        output_docx_path = 'design_original_file_name.docx'
        
        # Check if the file exists
        if not os.path.exists(output_docx_path):
            return Response({'error': 'File not found'}, status=status.HTTP_404_NOT_FOUND)

        # Return the file as a downloadable attachment
        response = FileResponse(open(output_docx_path, 'rb'), as_attachment=True)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(output_docx_path)}"'
        return response

# API View for downloading the replaced DOCX document
class DownloadReplacedDocxAPI(APIView):
    def get(self, request, *args, **kwargs):
        replaced_output_docx_path = 'replaced_design_original_file_name.docx'
        
        # Check if the file exists
        if not os.path.exists(replaced_output_docx_path):
            return Response({'error': 'File not found'}, status=status.HTTP_404_NOT_FOUND)

        # Return the file as a downloadable attachment
        response = FileResponse(open(replaced_output_docx_path, 'rb'), as_attachment=True)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(replaced_output_docx_path)}"'
        return response
